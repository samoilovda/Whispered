"""
Whispered UI - Export Functions
Export transcription results to various formats
"""

from transcriber import TranscriptionResult
from utils import format_timestamp_srt, format_timestamp_vtt


def _speaker_name(result: TranscriptionResult, seg) -> str:
    """Display name for a segment's speaker ('' if none).

    Honours user renames stored in result.speaker_names; falls back to the raw
    speaker id. Works even on results that predate the speaker_names field.
    """
    speaker_id = getattr(seg, 'speaker', None)
    if not speaker_id:
        return ''
    names = getattr(result, 'speaker_names', None) or {}
    return names.get(speaker_id, speaker_id)


def export_txt(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as plain text."""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(result.full_text)


def export_txt_with_timestamps(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as text with timestamps."""
    with open(filepath, 'w', encoding='utf-8') as f:
        for seg in result.segments:
            timestamp = format_timestamp_vtt(seg.start)
            speaker = _speaker_name(result, seg)
            prefix = f"{speaker}: " if speaker else ""
            f.write(f"[{timestamp}] {prefix}{seg.text.strip()}\n")


def export_srt(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as SRT subtitle file."""
    with open(filepath, 'w', encoding='utf-8') as f:
        for i, seg in enumerate(result.segments, start=1):
            start = format_timestamp_srt(seg.start)
            end = format_timestamp_srt(seg.end)
            speaker = _speaker_name(result, seg)
            prefix = f"{speaker}: " if speaker else ""
            text = seg.text.strip()

            f.write(f"{i}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"{prefix}{text}\n")
            f.write("\n")


def export_vtt(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as WebVTT subtitle file."""
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write("WEBVTT\n\n")

        for i, seg in enumerate(result.segments, start=1):
            start = format_timestamp_vtt(seg.start)
            end = format_timestamp_vtt(seg.end)
            speaker = _speaker_name(result, seg)
            text = seg.text.strip()
            if speaker:
                # WebVTT voice spans must not contain '<' or '>'.
                safe_speaker = speaker.replace("<", "").replace(">", "").replace("\n", " ").strip()
                cue = f"<v {safe_speaker}>{text}</v>" if safe_speaker else text
            else:
                cue = text

            f.write(f"{i}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"{cue}\n")
            f.write("\n")


def export_json(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as JSON."""
    import json

    data = {
        'language': result.language,
        'duration': result.duration,
        'text': result.full_text,
        'speaker_names': getattr(result, 'speaker_names', None) or {},
        'segments': [
            {
                'start': seg.start,
                'end': seg.end,
                'text': seg.text.strip(),
                'speaker_id': getattr(seg, 'speaker', None),
                'speaker': _speaker_name(result, seg) or None,
            }
            for seg in result.segments
        ]
    }

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def export_md(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as Markdown."""
    import datetime

    minutes = int(result.duration // 60)
    seconds = int(result.duration % 60)
    duration_str = f"{minutes}m {seconds}s"
    date_str = datetime.datetime.now().strftime('%Y-%m-%d')

    with open(filepath, 'w', encoding='utf-8') as f:
        # YAML frontmatter
        f.write('---\n')
        f.write(f'date: {date_str}\n')
        f.write(f'language: {result.language}\n')
        f.write(f'duration: {duration_str}\n')
        f.write('---\n\n')

        # Full transcript body
        f.write('## Transcript\n\n')
        f.write(result.full_text)
        f.write('\n\n')

        # Timestamped segments section
        f.write('## Segments\n\n')
        for seg in result.segments:
            start = format_timestamp_vtt(seg.start)
            name = _speaker_name(result, seg)
            speaker = f'**{name}**: ' if name else ''
            f.write(f'- `{start}` {speaker}{seg.text.strip()}\n')


def export_html(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as standalone HTML with embedded CSS."""
    import html as html_lib
    import datetime

    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    lines = []
    for seg in result.segments:
        start = format_timestamp_vtt(seg.start)
        name = _speaker_name(result, seg)
        text = html_lib.escape(seg.text.strip())
        if name:
            safe_name = html_lib.escape(name)
            lines.append(
                f'<p class="seg"><span class="ts">[{start}]</span> '
                f'<strong class="spk">{safe_name}:</strong> {text}</p>'
            )
        else:
            lines.append(
                f'<p class="seg"><span class="ts">[{start}]</span> {text}</p>'
            )

    body = "\n".join(lines)
    doc = f"""<!DOCTYPE html>
<html lang="{result.language}">
<head>
<meta charset="utf-8">
<meta name="date" content="{date_str}">
<style>
  body {{ font-family: sans-serif; max-width: 860px; margin: 2rem auto;
         background: #f9f9f9; color: #222; line-height: 1.6; }}
  h1 {{ font-size: 1.3rem; border-bottom: 1px solid #ddd; padding-bottom: .4rem; }}
  .seg {{ margin: .4rem 0; }}
  .ts {{ color: #888; font-size: .8em; }}
  .spk {{ color: #6366f1; }}
</style>
</head>
<body>
<h1>Transcript &mdash; {date_str}</h1>
{body}
</body>
</html>"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(doc)


def export_docx(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as DOCX (requires python-docx)."""
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed.\n\nInstall it with:\n  pip install python-docx"
        )
    import datetime

    doc = docx.Document()
    doc.add_heading("Transcript", level=1)

    # Metadata paragraph
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    minutes = int(result.duration // 60)
    seconds = int(result.duration % 60)
    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date_str}  |  Language: {result.language}  |  Duration: {minutes}m {seconds}s")
    meta.style = "Body Text"
    doc.add_paragraph()  # spacer

    for seg in result.segments:
        start = format_timestamp_vtt(seg.start)
        name = _speaker_name(result, seg)
        para = doc.add_paragraph()
        # Timestamp in grey
        ts_run = para.add_run(f"[{start}] ")
        ts_run.font.color.rgb = docx.shared.RGBColor(0x88, 0x88, 0x88)
        ts_run.font.size = docx.shared.Pt(9)
        # Speaker in bold
        if name:
            spk_run = para.add_run(f"{name}: ")
            spk_run.bold = True
        # Text
        para.add_run(seg.text.strip())

    doc.save(filepath)


def export_pdf(result: TranscriptionResult, filepath: str) -> None:
    """Export transcription as PDF via Qt (no extra dependencies).

    Must be called from the main Qt thread (or any thread with QApplication).
    """
    try:
        from PyQt6.QtGui import QTextDocument
        from PyQt6.QtPrintSupport import QPrinter
    except ImportError:
        raise RuntimeError(
            "PyQt6.QtPrintSupport is required for PDF export. "
            "Ensure the full PyQt6 package is installed."
        )

    # Build HTML and render via QTextDocument
    import io
    buf = io.StringIO()
    buf.write("<html><head><meta charset='utf-8'></head><body>")
    buf.write("<h2>Transcript</h2>")
    for seg in result.segments:
        from utils import format_timestamp_vtt
        start = format_timestamp_vtt(seg.start)
        name = _speaker_name(result, seg)
        text = seg.text.strip().replace("<", "&lt;").replace(">", "&gt;")
        prefix = f"<b>{name}:</b> " if name else ""
        buf.write(f"<p><small>[{start}]</small> {prefix}{text}</p>")
    buf.write("</body></html>")

    html = buf.getvalue()
    td = QTextDocument()
    td.setHtml(html)

    printer = QPrinter(QPrinter.PrinterMode.HighResolution)
    printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
    printer.setOutputFileName(filepath)
    td.print(printer)


# Export format options
EXPORT_FORMATS = {
    'txt': ('Plain Text (.txt)', export_txt),
    'txt_ts': ('Text with Timestamps (.txt)', export_txt_with_timestamps),
    'srt': ('SRT Subtitles (.srt)', export_srt),
    'vtt': ('WebVTT Subtitles (.vtt)', export_vtt),
    'json': ('JSON (.json)', export_json),
    'md': ('Markdown (.md)', export_md),
    'html': ('HTML (.html)', export_html),
    'docx': ('Word Document (.docx)', export_docx),
    'pdf': ('PDF (.pdf)', export_pdf),
}


def export_result(
    result: TranscriptionResult,
    filepath: str,
    format_key: str
) -> None:
    """
    Export transcription result to specified format.

    Args:
        result: The transcription result
        filepath: Output file path
        format_key: One of 'txt', 'txt_ts', 'srt', 'vtt', 'json'
    """
    if format_key not in EXPORT_FORMATS:
        raise ValueError(f"Unknown format: {format_key}")

    _, export_func = EXPORT_FORMATS[format_key]
    export_func(result, filepath)
