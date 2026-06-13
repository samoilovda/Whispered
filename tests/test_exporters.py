"""
Tests for exporters.py.
No Qt, no GPU, no whisper binary required.

We stub the transcriber module to avoid its Qt/pywhispercpp dependencies,
then import exporters which accesses TranscriptionResult via duck-typing.
"""
import json
import os
import sys
import types
from dataclasses import dataclass, field
from typing import Optional, List
import pytest


# ---------------------------------------------------------------------------
# Stub transcriber module (avoids PyQt6 / pywhispercpp at import time)
# ---------------------------------------------------------------------------

@dataclass
class _Segment:
    start: float
    end: float
    text: str
    speaker: Optional[str] = None


@dataclass
class _TranscriptionResult:
    segments: List[_Segment]
    language: str
    duration: float
    speaker_names: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return " ".join(s.text.strip() for s in self.segments)


# Inject stub before importing exporters
_stub_transcriber = types.ModuleType("transcriber")
_stub_transcriber.Segment = _Segment
_stub_transcriber.TranscriptionResult = _TranscriptionResult
sys.modules.setdefault("transcriber", _stub_transcriber)

# Also stub pywhispercpp so that any late import doesn't crash
sys.modules.setdefault("pywhispercpp", types.ModuleType("pywhispercpp"))
sys.modules.setdefault("pywhispercpp.model", types.ModuleType("pywhispercpp.model"))

from exporters import (   # noqa: E402  (after stub setup)
    export_txt,
    export_txt_with_timestamps,
    export_srt,
    export_vtt,
    export_json,
    export_md,
    export_html,
    export_docx,
    export_result,
    EXPORT_FORMATS,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_result(has_speakers: bool = False) -> _TranscriptionResult:
    segments = [
        _Segment(0.0,  5.0,  " Привет мир ",  "Speaker 1" if has_speakers else None),
        _Segment(5.0,  10.0, " Hello world ",  "Speaker 2" if has_speakers else None),
        _Segment(10.0, 15.0, " Тест юникод ", None),
    ]
    return _TranscriptionResult(segments=segments, language="ru", duration=15.0)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestExportTxt:
    def test_creates_file(self, tmp_path):
        out = str(tmp_path / "out.txt")
        export_txt(_make_result(), out)
        assert os.path.exists(out)

    def test_contains_cyrillic(self, tmp_path):
        out = str(tmp_path / "out.txt")
        export_txt(_make_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "Привет" in content

    def test_unicode_preserved(self, tmp_path):
        out = str(tmp_path / "out.txt")
        export_txt(_make_result(), out)
        assert "юникод" in open(out, encoding="utf-8").read()


class TestExportTxtWithTimestamps:
    def test_timestamps_present(self, tmp_path):
        out = str(tmp_path / "out_ts.txt")
        export_txt_with_timestamps(_make_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "[" in content


class TestExportSrt:
    def test_starts_with_index(self, tmp_path):
        out = str(tmp_path / "out.srt")
        export_srt(_make_result(), out)
        assert open(out, encoding="utf-8").read().startswith("1\n")

    def test_arrow_separator(self, tmp_path):
        out = str(tmp_path / "out.srt")
        export_srt(_make_result(), out)
        assert "-->" in open(out, encoding="utf-8").read()

    def test_three_entries(self, tmp_path):
        out = str(tmp_path / "out.srt")
        export_srt(_make_result(), out)
        entries = [b for b in open(out, encoding="utf-8").read().strip().split("\n\n") if b.strip()]
        assert len(entries) == 3

    def test_comma_in_timestamp(self, tmp_path):
        out = str(tmp_path / "out.srt")
        export_srt(_make_result(), out)
        assert "," in open(out, encoding="utf-8").read()


class TestExportVtt:
    def test_webvtt_header(self, tmp_path):
        out = str(tmp_path / "out.vtt")
        export_vtt(_make_result(), out)
        assert open(out, encoding="utf-8").read().startswith("WEBVTT")

    def test_dot_separator_in_timestamps(self, tmp_path):
        out = str(tmp_path / "out.vtt")
        export_vtt(_make_result(), out)
        assert "00:00:00.000" in open(out, encoding="utf-8").read()

    def test_voice_tag_special_chars_sanitized(self, tmp_path):
        """Speaker name with '<' or '>' must not corrupt the VTT cue."""
        out = str(tmp_path / "bad_speaker.vtt")
        r = _make_result(has_speakers=True)
        r.speaker_names = {"Speaker 1": "Bob > Alice", "Speaker 2": "Eve<2>"}
        export_vtt(r, out)
        content = open(out, encoding="utf-8").read()
        # The resulting voice tag must never contain a raw '>'  after <v
        for line in content.splitlines():
            if line.startswith("<v "):
                # Everything between <v and the first closing > is the name
                # The name itself must not contain '>'.
                name_part = line[3:line.index(">")]
                assert ">" not in name_part


class TestExportJson:
    def test_valid_json(self, tmp_path):
        out = str(tmp_path / "out.json")
        export_json(_make_result(), out)
        data = json.load(open(out, encoding="utf-8"))
        assert "text" in data
        assert "segments" in data
        assert "language" in data

    def test_segment_fields(self, tmp_path):
        out = str(tmp_path / "out.json")
        export_json(_make_result(), out)
        seg = json.load(open(out, encoding="utf-8"))["segments"][0]
        assert {"start", "end", "text"} <= set(seg)

    def test_unicode_not_escaped(self, tmp_path):
        out = str(tmp_path / "out.json")
        export_json(_make_result(), out)
        raw = open(out, encoding="utf-8").read()
        assert "Привет" in raw  # ensure_ascii=False


class TestExportMd:
    def test_yaml_frontmatter(self, tmp_path):
        out = str(tmp_path / "out.md")
        export_md(_make_result(), out)
        content = open(out, encoding="utf-8").read()
        assert content.startswith("---")
        assert "language:" in content

    def test_transcript_section(self, tmp_path):
        out = str(tmp_path / "out.md")
        export_md(_make_result(), out)
        assert "## Transcript" in open(out, encoding="utf-8").read()

    def test_segments_section(self, tmp_path):
        out = str(tmp_path / "out.md")
        export_md(_make_result(), out)
        assert "## Segments" in open(out, encoding="utf-8").read()

    def test_speaker_in_output(self, tmp_path):
        out = str(tmp_path / "out.md")
        export_md(_make_result(has_speakers=True), out)
        assert "Speaker" in open(out, encoding="utf-8").read()


def _make_named_result() -> _TranscriptionResult:
    """Result with speakers renamed via speaker_names."""
    r = _make_result(has_speakers=True)
    r.speaker_names = {"Speaker 1": "Alice", "Speaker 2": "Bob"}
    return r


class TestSpeakerNamesInExport:
    def test_md_uses_display_name(self, tmp_path):
        out = str(tmp_path / "out.md")
        export_md(_make_named_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "Alice" in content and "Bob" in content
        assert "Speaker 1" not in content

    def test_txt_ts_uses_display_name(self, tmp_path):
        out = str(tmp_path / "out_ts.txt")
        export_txt_with_timestamps(_make_named_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "Alice:" in content
        assert "Speaker 1" not in content

    def test_srt_uses_display_name(self, tmp_path):
        out = str(tmp_path / "out.srt")
        export_srt(_make_named_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "Alice:" in content

    def test_vtt_voice_tag(self, tmp_path):
        out = str(tmp_path / "out.vtt")
        export_vtt(_make_named_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "<v Alice>" in content

    def test_json_includes_speaker_fields(self, tmp_path):
        out = str(tmp_path / "out.json")
        export_json(_make_named_result(), out)
        data = json.load(open(out, encoding="utf-8"))
        assert data["speaker_names"] == {"Speaker 1": "Alice", "Speaker 2": "Bob"}
        seg0 = data["segments"][0]
        assert seg0["speaker"] == "Alice"
        assert seg0["speaker_id"] == "Speaker 1"

    def test_no_speakers_no_prefix(self, tmp_path):
        out = str(tmp_path / "out_ts.txt")
        export_txt_with_timestamps(_make_result(has_speakers=False), out)
        content = open(out, encoding="utf-8").read()
        assert ":" not in content.split("]")[1].split("\n")[0]  # no "Name:" prefix


class TestExportHtml:
    def test_creates_html_file(self, tmp_path):
        out = str(tmp_path / "out.html")
        export_html(_make_result(), out)
        assert os.path.exists(out)

    def test_html_contains_doctype(self, tmp_path):
        out = str(tmp_path / "out.html")
        export_html(_make_result(), out)
        assert "<!DOCTYPE html>" in open(out, encoding="utf-8").read()

    def test_html_contains_cyrillic(self, tmp_path):
        out = str(tmp_path / "out.html")
        export_html(_make_result(), out)
        assert "Привет" in open(out, encoding="utf-8").read()

    def test_html_speaker_bold(self, tmp_path):
        out = str(tmp_path / "out.html")
        export_html(_make_named_result(), out)
        content = open(out, encoding="utf-8").read()
        assert "<strong" in content
        assert "Alice" in content


class TestExportDocx:
    def test_creates_docx_file(self, tmp_path):
        out = str(tmp_path / "out.docx")
        export_docx(_make_result(), out)
        assert os.path.exists(out)
        assert os.path.getsize(out) > 0

    def test_docx_readable(self, tmp_path):
        pytest.importorskip("docx")
        import docx as docx_lib
        out = str(tmp_path / "out.docx")
        export_docx(_make_result(), out)
        doc = docx_lib.Document(out)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Привет" in full_text

    def test_docx_speaker_names(self, tmp_path):
        pytest.importorskip("docx")
        import docx as docx_lib
        out = str(tmp_path / "out.docx")
        export_docx(_make_named_result(), out)
        doc = docx_lib.Document(out)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Alice" in full_text


class TestExportResult:
    def test_all_format_keys(self, tmp_path):
        result = _make_result()
        # PDF requires Qt display — skip in headless CI
        skip_formats = {"pdf"}
        for key in EXPORT_FORMATS:
            if key in skip_formats:
                continue
            ext = "txt" if key in ("txt", "txt_ts") else key
            out = str(tmp_path / f"out_{key}.{ext}")
            export_result(result, out, key)
            assert os.path.exists(out), f"Missing output for {key}"
            assert os.path.getsize(out) > 0, f"Empty output for {key}"

    def test_unknown_format_raises(self, tmp_path):
        with pytest.raises((ValueError, KeyError)):
            export_result(_make_result(), str(tmp_path / "x.xyz"), "unknown_fmt")
